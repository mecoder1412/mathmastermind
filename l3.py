import streamlit as st
import config
import io
import google.genai as genai
import time
from docx import Document
from docx.shared import Inches
# Configure your API key
client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])

def generate_response(prompt: str) -> str:
    # ... (system_prompt remains the same) ...
    system_prompt = """You are a Math Mastermind - an expert mathematics problem solver with exceptional abilities in:
- Algebra, Calculus, Geometry, Trigonometry
- Statistics, Probability, Linear Algebra
- Discrete Mathematics, Number Theory
- Mathematical Proofs and Logic
- Applied Mathematics and Word Problems

For every math problem:
1. Show clear step-by-step solutions
2. Explain the mathematical reasoning
3. Provide alternative solving methods when applicable
4. Verify your answer when possible
5. Use proper mathematical notation
6. Break down complex problems into manageable parts

Format your responses with:
- Clear problem identification
- Step-by-step solution process
- Final answer highlighted
- Brief explanation of concepts used

Always be precise, thorough, and educational in your mathematical explanations."""

    # <--- START: ADDED RETRY LOGIC HERE
    max_retries = 3
    for attempt in range(max_retries):
        try:
            # CORRECT SDK syntax for system prompt and single-turn content
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=[
                    {"role": "user", "parts": [{"text": system_prompt}]},
                    {"role": "user", "parts": [{"text": prompt}]}
                ],
                config={
                    "temperature": 0.1,
                }
            )

            # If successful, return the response text immediately
            return response.text or "Error: No response from model"

        except Exception as e:
            error_message = str(e)
            # Check if the error is specifically the 503 UNAVAILABLE error
            if "503 UNAVAILABLE" in error_message and attempt < max_retries - 1:
                # Wait for a short period before retrying (exponential backoff)
                wait_time = 2 ** attempt  # waits 1s, then 2s
                print(f"Service overloaded. Retrying in {wait_time} seconds... (Attempt {attempt + 1}/{max_retries})")
                time.sleep(wait_time)
            else:
                # If it's a different error or we're on the last attempt, return the error
                return f"Error: {error_message}"
    
    # Return an error message if all retries fail
    return "Error: Failed after multiple retries due to service unavailability."
    # <--- END: ADDED RETRY LOGIC HERE


def setup_ui():
# ... (rest of your setup_ui function remains exactly the same) ...
# ... (rest of your setup_ui function remains exactly the same) ...
# ... (rest of your setup_ui function remains exactly the same) ...
    st.set_page_config(page_title="🧮 Math Mastermind", layout="centered")
    st.title("🧮 Math Mastermind")
    st.write("**Your Expert Mathematical Problem Solver** - From basic arithmetic to advanced calculus, I'll solve any math problem with detailed step-by-step explanations!")

    # Examples
    with st.expander("📚 Example Problems I Can Solve"):
        st.markdown("""
        **Algebra:** Solve equations, factor polynomials, simplify expressions
        - Example: "Solve 2x² + 5x - 3 = 0"

        **Calculus:** Derivatives, integrals, limits, optimization
        - Example: "Find the derivative of sin(x²) + ln(x)"

        **Geometry:** Area, volume, proofs, coordinate geometry
        - Example: "Find the area of a triangle with vertices at (0,0), (3,4), and (6,0)"

        **Statistics:** Probability, distributions, hypothesis testing
        - Example: "What's the probability of rolling two dice and getting a sum of 7?"

        **Word Problems:** Real-world applications of mathematics
        - Example: "A train travels 300 miles in 4 hours. How fast was it going?"
        """)

    # Initialize session state
    if "history" not in st.session_state:
        st.session_state.history = []
    if "input_key" not in st.session_state:
        st.session_state.input_key = 0

    # Clear and Export buttons
    col_clear, col_export = st.columns([1, 2])
    with col_clear:
        if st.button("🧹 Clear Conversation"):
            st.session_state.history = []
            st.rerun()
    with col_export:
        if st.session_state.history:
            # --- START OF DOCX EXPORT LOGIC ---
            
            # Create a new Word document object
            document = Document()
            document.add_heading('Math Mastermind Solutions History', 0)
            
            total_questions = len(st.session_state.history)

            # Iterate through the history and add content to the document
            for idx, qa in enumerate(st.session_state.history):
                question_num = total_questions - idx
                
                # Add Question/Problem
                document.add_heading(f"Problem {question_num} ({qa.get('difficulty', 'N/A')})", level=1)
                document.add_paragraph(f"Question: {qa['question']}", style='List Bullet')
                
                # Add Answer/Solution
                document.add_heading(f"Solution {question_num}", level=2)
                
                # We need to preserve the formatting (like line breaks) from the answer
                # by splitting the text into paragraphs.
                solution_lines = qa['answer'].split('\n')
                for line in solution_lines:
                    if line.strip(): # Avoid adding empty lines as paragraphs
                        document.add_paragraph(line)
                    
                document.add_page_break() # Separate solutions with a page break for clarity

            # Save the document to an in-memory stream
            doc_io = io.BytesIO()
            document.save(doc_io)
            doc_io.seek(0)
            
            # --- END OF DOCX EXPORT LOGIC ---

            st.download_button(
                label="📥 Export Word Document (.docx)",
                data=doc_io, # Use the in-memory Word document stream
                file_name="Math_Mastermind_Solutions.docx", # Change file extension
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", # Change MIME type
            )
    # Input form
    with st.form(key="math_form", clear_on_submit=True):
        user_input = st.text_area(
            "🔢 Enter your math problem here:",
            height=100,
            placeholder="Example: Solve x² + 5x + 6 = 0 or Find the integral of 2x + 3",
            key=f"user_input_{st.session_state.input_key}"
        )

        col1, col2 = st.columns([3, 1])
        with col1:
            submitted = st.form_submit_button("🧮 Solve Problem", use_container_width=True)
        with col2:
            difficulty = st.selectbox("Level", ["Basic", "Intermediate", "Advanced"], index=1)

        if submitted and user_input.strip():
            enhanced_prompt = f"[{difficulty} Level] {user_input.strip()}"
            with st.spinner("🔍 Analyzing and solving your math problem..."):
                response = generate_response(enhanced_prompt)

            st.session_state.history.insert(0, {
                "question": user_input.strip(),
                "answer": response,
                "difficulty": difficulty
            })
            st.session_state.input_key += 1
            st.rerun()
        elif submitted:
            st.warning("⚠️ Please enter a math problem before clicking Solve Problem.")

    # Display history
    if st.session_state.history:
        st.markdown("### 📋 Solution History (Latest First)")
        st.markdown("""
        <style>
        .history-box { max-height: 500px; overflow-y: auto; border: 2px solid #4CAF50; padding: 15px; background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%); border-radius: 10px; font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);}
        .question { font-weight: 700; color: #2E7D32; margin-top: 15px; margin-bottom: 8px; font-size: 16px;}
        .difficulty { display: inline-block; background-color: #FF9800; color: white; padding: 2px 8px; border-radius: 12px; font-size: 12px; font-weight: bold; margin-left: 10px;}
        .answer { margin-bottom: 20px; white-space: pre-wrap; color: #1B5E20; line-height: 1.6; background-color: rgba(255,255,255,0.7); padding: 12px; border-radius: 8px; border-left: 4px solid #4CAF50;}
        </style>
        """, unsafe_allow_html=True)

        history_html = '<div class="history-box">'
        total_questions = len(st.session_state.history)
        for idx, qa in enumerate(st.session_state.history):
            question_num = total_questions - idx
            difficulty_badge = f'<span class="difficulty">{qa.get("difficulty", "N/A")}</span>' if "difficulty" in qa else ""
            history_html += f'<div class="question">Problem {question_num}: {qa["question"]}{difficulty_badge}</div>'
            history_html += f'<div class="answer">Solution {question_num}: {qa["answer"]}</div>'
        history_html += '</div>'
        st.markdown(history_html, unsafe_allow_html=True)


def main():
    setup_ui()


if __name__ == "__main__":
    main()